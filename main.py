# -*- coding: utf-8 -*-
import os
import json
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from screen import make_screens
from word_grafik import make_grafik


# ====== CONFIG (–ª—É—á—à–µ –≤—ã–Ω–µ—Å—Ç–∏ –≤ .env, –Ω–æ –º–æ–∂–Ω–æ —Ç–∞–∫) ======
PIC_W = Inches(6.5)
FONT_NAME = "Times New Roman"
FONT_SIZE = 14

UZ_MONTHS = {
    1: "yanvar", 2: "fevral", 3: "mart", 4: "aprel", 5: "may", 6: "iyun",
    7: "iyul", 8: "avgust", 9: "sentabr", 10: "oktabr", 11: "noyabr", 12: "dekabr"
}


def _uz_date(date_str: str) -> str:
    d = datetime.strptime(date_str, "%Y-%m-%d")
    return f"{d.year}-yil {d.day}-{UZ_MONTHS[d.month]}"


def _load_json(path: str) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _set_default_style(doc: Document):
    st = doc.styles["Normal"]
    st.font.name = FONT_NAME
    st.font.size = Pt(FONT_SIZE)


def _p_center_bold(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT_NAME
    r.font.size = Pt(FONT_SIZE)
    return p


def _p_left(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name = FONT_NAME
        r.font.size = Pt(FONT_SIZE)
    return p

def _p_justify(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # üîπ –ø–æ —à–∏—Ä–∏–Ω–µ
    for r in p.runs:
        r.font.name = FONT_NAME
        r.font.size = Pt(FONT_SIZE)
    return p

def _add_picture_center(doc: Document, path: str, width=PIC_W):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=width)
    return p


def _add_multiline(doc: Document, text: str):
    for part in [t.strip() for t in (text or "").split("\n\n") if t.strip()]:
        _p_justify(doc, part)   # ‚úÖ —Ç–æ–ª—å–∫–æ JSON ‚Äî –ø–æ —à–∏—Ä–∏–Ω–µ



def build_docx(
    gases,
    date_str: str,
    parent_cod: int,
    count_gase: int,
    *,
    rasters_root: str,
    mintaqa_shp: str,
    tuman_shp: str,
    text_json_path: str,
    out_docx: str,
):
    """
    –î–µ–ª–∞–µ—Ç docx. –í–Ω—É—Ç—Ä–∏ –≥–µ–Ω–µ—Ä–∏—Ç 3 PNG –Ω–∞ –≥–∞–∑:
      1) mintaqa_screen (–≤—Å—è —Ä–µ—Å–ø—É–±–ª–∏–∫–∞)
      2) grafik
      3) rayon_screen (—Ä–µ–≥–∏–æ–Ω + —Å–æ—Å–µ–¥–Ω–∏–µ)
    """

    # –Ω–æ—Ä–º–∞–ª–∏–∑—É–µ–º –≤—Ö–æ–¥
    gases = [g.strip().upper() for g in gases if str(g).strip()]
    if not gases:
        raise ValueError("gases is empty")

    if count_gase not in (7, 15, 30):
        count_gase = 30

    gas_db = _load_json(text_json_path)

    os.makedirs(os.path.dirname(out_docx), exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="airrep_") as tmpdir:
        doc = Document()
        _set_default_style(doc)

        for idx, gas in enumerate(gases, start=1):
            info = gas_db.get(gas, {})
            gas_name = info.get("display_uz", gas)
            unit = info.get("unit", "")
            text_tpl = info.get("text_uz", "")

            # 1) Screens (2 png)
            screens = make_screens(
                gas=gas,
                date_str=date_str,
                parent_cod=parent_cod,
                rasters_root=rasters_root,
                vector_path=mintaqa_shp,         # <-- –æ–¥–∏–Ω –∏ —Ç–æ—Ç –∂–µ shp
                base_vector_path=tuman_shp,
                out_dir=tmpdir,
            )

            # 2) Grafik (1 png)
            grafik = make_grafik(
                gas=gas,
                date_str=date_str,
                parent_cod=parent_cod,
                rasters_root=rasters_root,
                mintaqa_shp=mintaqa_shp,         # <-- —Ç–æ—Ç –∂–µ shp
                out_dir=tmpdir,
                lookback_days=count_gase,
            )
            region_name = grafik.get("region_name") or str(parent_cod)

            # ===== PAGE 1: —Ä–µ—Å–ø—É–±–ª–∏–∫–∞ png -> –≥—Ä–∞—Ñ–∏–∫ =====
            _p_center_bold(
                doc,
                f"{idx}. Respublika va {region_name} kesimida {_uz_date(date_str)} holatiga ko‚Äòra "
                f"{gas_name} yuqori bo‚Äòlgan hududlar."
            )

            _p_left(doc, "Respublika kesimida:")
            _add_picture_center(doc, screens["mintaqa"])

            _p_left(doc, f"So‚Äònggi {count_gase} kun bo‚Äòyicha {gas} o‚Äòrtacha qiymat grafigi:")
            _add_picture_center(doc, grafik["png"])

            doc.add_page_break()

            # ===== PAGE 2: region png -> text =====
            _p_left(doc, f"{region_name} va unga yondosh hududlar:")
            _add_picture_center(doc, screens["rayon"])

            if text_tpl:
                _add_multiline(doc, text_tpl.format(unit=unit))

            if idx != len(gases):
                doc.add_page_break()

        doc.save(out_docx)

    return out_docx


# ====== local test ======
if __name__ == "__main__":
    # –≤—Ö–æ–¥—ã –¥–ª—è –ª–æ–∫–∞–ª—å–Ω–æ–≥–æ –∑–∞–ø—É—Å–∫–∞
    GASES = ["AERAI", "CH4", "CO", "HCHO","O3", "SO2", "NO2"]
    # GASES = ["AERAI","CO",]
    DATE = "2025-12-27"
    PARENT_COD = 1726
    COUNT_GASE = 7


    RASTERS_ROOT = r"D:\Xalim\wind_visual\gaz\sentinel\data"
    MINTAQA_SHP = r"D:\Xalim\documents\Mintaqa\Mintaqa.shp"
    TUMAN_SHP = r"D:\Xalim\documents\Mintaqa\Tuman.shp"

    TEXT_JSON_PATH = r"D:\Xalim\wind_visual\text.json"
    OUT_DOCX = r"D:\Xalim\wind_visual\word\report12345.docx"

    build_docx(
        GASES, DATE, PARENT_COD, COUNT_GASE,
        rasters_root=RASTERS_ROOT,
        mintaqa_shp=MINTAQA_SHP,
        tuman_shp=TUMAN_SHP,
        text_json_path=TEXT_JSON_PATH,
        out_docx=OUT_DOCX,
    )
    print("Saved:", OUT_DOCX)
